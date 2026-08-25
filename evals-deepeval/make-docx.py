"""
Builds the Word version of the evaluation report.

python-docx has no public hyperlink API, so add_hyperlink() writes the
relationship and the w:hyperlink element directly - that is what makes the
~24 reference links actually clickable in Word rather than plain blue text.

  evals-deepeval/.venv/Scripts/python.exe evals-deepeval/make-docx.py
"""

import pathlib

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

OUT = pathlib.Path(__file__).resolve().parent / "Chatbot-Quality-Evaluation.docx"

# palette mirrors the HTML report
INK = RGBColor(0x12, 0x23, 0x2B)
ACCENT = RGBColor(0x0F, 0x6E, 0x6E)
MUTED = RGBColor(0x5C, 0x6F, 0x75)
GOOD = RGBColor(0x2F, 0x7D, 0x5B)
POOR = RGBColor(0xA6, 0x3D, 0x40)

SANS = "IBM Plex Sans"
SERIF = "IBM Plex Serif"
MONO = "IBM Plex Mono"
# fallbacks Word will find on any machine
SANS_FB, SERIF_FB, MONO_FB = "Segoe UI", "Georgia", "Consolas"


def set_font(run, name, size=None, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size) if size else run.font.size
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    # east-asian + complex-script fallback so Word doesn't substitute silently
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), name)


def add_hyperlink(paragraph, url, text, size=10.5):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")

    fonts = OxmlElement("w:rFonts")
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        fonts.set(qn(attr), SANS)
    rpr.append(fonts)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0F6E6E")
    rpr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    rpr.append(sz)

    run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    link.append(run)
    paragraph._p.append(link)
    return link


def shade(cell, hex_color):
    tcpr = cell._tc.get_or_add_tcPr()
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hex_color)
    tcpr.append(el)


def para(doc, text="", font=SERIF, size=10.5, bold=False, italic=False,
         color=INK, space_after=8, space_before=0, align=None, indent=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if align:
        p.alignment = align
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    if text:
        set_font(p.add_run(text), font, size, bold, italic, color)
    return p


def rich(doc, parts, size=10.5, space_after=8, indent=None):
    """parts = [(text, bold, font)]"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    for text, bold, font in parts:
        set_font(p.add_run(text), font, size, bold, False, INK)
    return p


def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after = Pt(10)
    set_font(p.add_run(text), SANS, 15, True, False, INK)
    # rule under section headings
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), "0F6E6E")
    pbdr.append(bottom)
    p._p.get_or_add_pPr().append(pbdr)
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(5)
    set_font(p.add_run(text), SANS, 11.5, True, False, INK)
    return p


def bullet(doc, parts, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    for text, bold, font in parts:
        set_font(p.add_run(text), font, size, bold, False, INK)
    return p


def numbered(doc, parts, size=10.5):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    for text, bold, font in parts:
        set_font(p.add_run(text), font, size, bold, False, INK)
    return p


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        set_font(p.add_run(htext.upper()), MONO, 8, True, False, MUTED)
        shade(hdr[i], "EAF1F1")
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            mono = isinstance(val, str) and val.replace(".", "").replace("/", "").replace("~", "").replace("$", "").replace(",", "").replace("k", "").replace(" ", "").isdigit()
            colour = INK
            if isinstance(val, str) and val.startswith("!"):
                val, colour = val[1:], POOR
                mono = True
            elif isinstance(val, str) and val.startswith("+"):
                val, colour = val[1:], GOOD
                mono = True
            set_font(p.add_run(str(val)), MONO if mono else SANS, 9.5, False, False, colour)
    if widths:
        for i, w in enumerate(widths):
            for r in t.rows:
                r.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def qa(doc, who, exchanges, tone="bad"):
    """Question/answer exchange in a single-cell shaded box."""
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    cell = t.rows[0].cells[0]
    cell.text = ""
    shade(cell, "F4F8F8")
    accent = POOR if tone == "bad" else (GOOD if tone == "good" else ACCENT)

    first = True
    if who:
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(6)
        set_font(p.add_run(who.upper()), MONO, 8, True, False, MUTED)
        first = False

    for label, text in exchanges:
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.paragraph_format.space_after = Pt(5)
        set_font(p.add_run(f"{label}  "), MONO, 9.5, True, False, accent)
        set_font(p.add_run(text), SERIF, 10.5, False, False, INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def note(doc, parts):
    t = doc.add_table(rows=1, cols=1)
    cell = t.rows[0].cells[0]
    cell.text = ""
    shade(cell, "EAF1F1")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    for text, bold in parts:
        set_font(p.add_run(text), SERIF, 10.5, bold, False, INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def code(doc, text):
    t = doc.add_table(rows=1, cols=1)
    cell = t.rows[0].cells[0]
    cell.text = ""
    shade(cell, "F1F5F5")
    first = True
    for line in text.split("\n"):
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.paragraph_format.space_after = Pt(0)
        set_font(p.add_run(line), MONO, 9, False, False, INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def links(doc, items):
    for url, text, desc in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Inches(0.16)
        add_hyperlink(p, url, text)
        if desc:
            set_font(p.add_run(f"  — {desc}"), SANS, 9.5, False, False, MUTED)


# ══════════════════════════════════════════════════════════════════
doc = Document()

sec = doc.sections[0]
sec.left_margin = sec.right_margin = Inches(1.05)
sec.top_margin = Inches(0.95)
sec.bottom_margin = Inches(0.95)

normal = doc.styles["Normal"]
normal.font.name = SERIF
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.22

# ── cover ─────────────────────────────────────────────────────────
para(doc, "ENGINEERING REPORT", MONO, 9, True, False, ACCENT, space_after=10)
para(doc, "Chatbot Quality Evaluation", SANS, 26, True, False, INK, space_after=6)
para(doc, "Approach, tooling and measured results", SERIF, 13, False, True, MUTED, space_after=16)

table(doc, ["Field", "Detail"], [
    ["Date", "20 August 2026"],
    ["Scope", "RAG retrieval quality, answer grounding, tool selection, multi-turn behaviour"],
    ["Method", "DeepEval, run against live indexed customer content"],
    ["Environment", "Local, against the running application"],
], widths=[1.3, 4.9])

# ── 1 ─────────────────────────────────────────────────────────────
h1(doc, "1.  Summary")
para(doc, "Manual testing found that chatbots were sometimes answering from the language model's own general knowledge rather than the customer's uploaded content. To measure the problem rather than spot-check it, we introduced automated evaluation using DeepEval, an open-source LLM evaluation framework.")

h2(doc, "Headline results after fixes")
table(doc, ["Layer", "Metric", "Score"], [
    ["Component", "Contextual Recall", "+0.969"],
    ["Component", "Contextual Precision", "+0.956"],
    ["Component", "Contextual Relevancy", "!0.622"],
    ["RAG pipeline", "Faithfulness", "+0.955"],
    ["RAG pipeline", "Answer Relevancy", "+0.946"],
    ["Application", "Role Adherence", "+0.980"],
    ["Safety", "Toxicity / Bias", "+0.000 / 0.037"],
], widths=[1.5, 3.0, 1.7])
note(doc, [("The clearest remaining issue is content ingestion. ", True),
           ("Crawled web pages score 0.467 on relevancy versus 0.845 for clean PDF documents.", False)])

# ── 2 ─────────────────────────────────────────────────────────────
h1(doc, "2.  Manual testing and what it found")
para(doc, "We tested chatbots by hand across several assistant types, checking whether answers came from the customer's own content. Three problems surfaced.")

h2(doc, "Finding 1 — Answers from general knowledge, not customer data")
qa(doc, "IT-policy chatbot", [
    ("Q", "How do I reset my Netflix password?"),
    ("A", "To reset your Netflix password, please follow these steps: Visit the Netflix Password Recovery page. Choose how you want to reset your password: via email or text message. Enter your registered email address…"),
])
qa(doc, "Float-therapy chatbot", [
    ("Q", "Who won the world cup in 2018?"),
    ("A", "France won the FIFA World Cup in 2018. They defeated Croatia 4-2 in the final."),
])
para(doc, "Both answered confidently. Neither question relates to the business.")
rich(doc, [("A third case was subtler: a chatbot whose indexed content was a climate document answered a question ", False, SERIF),
           ("about that very content", True, SERIF),
           (" from memory, without searching. The answer was correct, so a human reviewer would have passed it — but it never touched the customer's data.", False, SERIF)])

h2(doc, "Finding 2 — Wrong tool selected for a core question")
qa(doc, None, [
    ("Q", "How much is a float session?"),
    ("A", "Routed to an internal get_services function that had no implementation, returning “This functionality will be available soon” — 8 times out of 8."),
])
para(doc, "Pricing is among the most common customer questions.")

h2(doc, "Finding 3 — Claiming actions it cannot perform")
qa(doc, None, [
    ("Q", "Can I book one for tomorrow at 6pm?"),
    ("A", "Great! I'll book you a session for tomorrow, August 20th, at 6 PM. Could you please provide your name, email, and phone number?"),
])
para(doc, "No booking system was connected. Nothing was booked, and personal details were collected for no purpose.")

h2(doc, "Why manual testing was not enough")
para(doc, "It does not scale, and it cannot measure. A fluent, plausible answer looks correct to a human reviewer even when it is ungrounded or self-contradictory.")

# ── 3 ─────────────────────────────────────────────────────────────
h1(doc, "3.  Why we moved to layered, measured evaluation")
para(doc, "We split evaluation into three layers so a failure can be located rather than guessed at.")
table(doc, ["Layer", "Question it answers"], [
    ["Component", "Does the retriever find the right content?"],
    ["RAG pipeline", "Is the generated answer grounded in what was retrieved?"],
    ["Application", "Does the assistant behave correctly across a whole conversation?"],
], widths=[1.5, 4.7])
para(doc, "Without this separation, a bad answer could be caused by ingestion, retrieval, ranking, the prompt, or the model — with no way to tell which.")

# ── 4 ─────────────────────────────────────────────────────────────
h1(doc, "4.  What is DeepEval")
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(8)
add_hyperlink(p, "https://deepeval.com/docs", "DeepEval")
set_font(p.add_run(" is an open-source framework for evaluating LLM applications — essentially Pytest, but specialised for testing model outputs."), SERIF, 10.5, False, False, INK)

table(doc, ["Attribute", "Value"], [
    ["Maintainer", "Confident AI"],
    ["Licence", "Apache 2.0"],
    ["GitHub", "~17.5k stars"],
    ["Languages", "Python (reference implementation), TypeScript (beta)"],
    ["Install", "pip install -U deepeval  ·  npm install --save-dev deepeval"],
    ["Version used", "4.1.8 (Python)"],
], widths=[1.5, 4.7])

h2(doc, "Core concepts")
bullet(doc, [("Test case", True, SANS), (" — a single interaction, holding the input, the actual output, the retrieved context and the tools called.", False, SERIF)])
bullet(doc, [("Golden", True, SANS), (" — a test case before execution: the question and expected answer, with no actual output. The application fills that in at run time, which is what makes one set reusable across model versions.", False, SERIF)])
bullet(doc, [("Dataset", True, SANS), (" — a collection of goldens, loadable from JSON, CSV or JSONL.", False, SERIF)])
bullet(doc, [("Metric", True, SANS), (" — a scorer producing 0–1 against a pass threshold. Most are LLM-as-a-judge.", False, SERIF)])
bullet(doc, [("Synthesizer", True, SANS), (" — generates goldens from your own documents or content.", False, SERIF)])
bullet(doc, [("ConversationSimulator", True, SANS), (" — plays the user across multiple turns so an application can be tested conversationally.", False, SERIF)])
para(doc, "It runs entirely locally; the Confident AI cloud platform is optional.")

h2(doc, "Metric families used")
table(doc, ["Family", "Metrics"], [
    ["RAG", "Faithfulness, Answer Relevancy, Contextual Precision, Contextual Recall, Contextual Relevancy"],
    ["Conversational", "Knowledge Retention, Conversation Completeness, Role Adherence, Topic Adherence"],
    ["Safety", "Bias, Toxicity, PII Leakage"],
    ["Agentic", "Tool Correctness, Argument Correctness, Task Completion, Step Efficiency"],
], widths=[1.3, 4.9])

# ── 5 ─────────────────────────────────────────────────────────────
h1(doc, "5.  Why we selected it, and its limitations")
h2(doc, "Why")
bullet(doc, [("It provides the exact metrics this problem needs — particularly Faithfulness (is every claim supported by retrieved context?) and Contextual Relevancy (how much of what we retrieved was useful?).", False, SERIF)])
bullet(doc, [("It can generate test questions from our own indexed content, so tests are grounded in real customer data rather than invented examples.", False, SERIF)])
bullet(doc, [("It separates component, RAG and application layers, matching our approach.", False, SERIF)])
bullet(doc, [("Open source, Apache 2.0, runs locally, no vendor lock-in.", False, SERIF)])

h2(doc, "Limitations we hit, recorded honestly")
numbered(doc, [("Golden generation is Python-only. ", True, SANS), ("The TypeScript package has no synthesizer, so generation and scoring run in Python while execution runs against the live TypeScript app.", False, SERIF)])
numbered(doc, [("A crash in its caching layer ", True, SANS), ("at batch size 28. Worked around by disabling cache; it only affects repeat-run speed.", False, SERIF)])
numbered(doc, [("Judge errors occur. ", True, SANS), ("One case flagged a business's public street address as a PII leak, while the judge's own explanation conceded it was public and not a privacy risk.", False, SERIF)])
numbered(doc, [("Some metrics penalise correct behaviour ", True, SANS), ("— see section 10. Judge explanations must be read, not just thresholded.", False, SERIF)])
numbered(doc, [("Telemetry is on by default. ", True, SANS), ("Disabled via an environment flag.", False, SERIF)])

# ── 6 ─────────────────────────────────────────────────────────────
h1(doc, "6.  What data we gave it")
rich(doc, [("We did ", False, SERIF), ("not", True, SERIF),
           (" hand-write questions or use sample data. DeepEval was given the customers' own indexed content — the exact text chunks stored in the vector database that the chatbot searches at runtime.", False, SERIF)])

h2(doc, "Chatbots covered")
table(doc, ["Chatbot", "Content source", "Format", "Assistant type"], [
    ["Float-therapy spa", "floatco.com", "Crawled website", "Customer support"],
    ["IT policy", "corporate IT policy", "PDF", "IT support"],
    ["Hong Kong food store", "store description + Shopify", "PDF", "E-commerce"],
    ["Developer guide", "technical guide", "Word document", "Sales"],
], widths=[1.6, 1.9, 1.4, 1.3])
para(doc, "Four content formats and four assistant types, so results are not specific to one customer or one ingestion path.")

h2(doc, "How content was sampled")
para(doc, "For each chatbot we ran four different topic searches against its vector database, using the returned chunks as generation context. For the spa:")
code(doc, """1. pricing membership packages cost per session
2. opening hours location address parking how to get there
3. what happens during a session, what to bring, preparation
4. policies, age limits, pregnancy, health restrictions, hygiene""")
note(doc, [("Why four and not one: ", True),
           ("an initial attempt with a single search produced four questions that were all about pricing, because the top-ranked chunks were all pricing pages. Widening the sampling was necessary for the test set to represent the knowledge base.", False)])

h2(doc, "What was generated")
bullet(doc, [("30 single-turn questions", True, SANS), (", each with a model answer and the source chunks", False, SERIF)])
bullet(doc, [("10 multi-turn conversation scenarios", True, SANS), ("", False, SERIF)])
para(doc, "Questions are deliberately harder than hand-written ones, using DeepEval's evolutions — reasoning, multi-context, comparative, hypothetical:")
qa(doc, None, [
    ("Q", "How do backup frequency, backup type, and recovery time objective (RTO) affect how backups are scheduled and stored?"),
    ("Q", "Can you tell me about the different float therapy membership tiers, their pricing, any weekday discounts, and the billing cycle?"),
], tone="neutral")

# ── 7 ─────────────────────────────────────────────────────────────
h1(doc, "7.  How the evaluation runs")
rich(doc, [("Three stages. ", False, SERIF), ("DeepEval does not run the application", True, SERIF),
           (" — it generates tests and scores results; execution hits our real endpoints.", False, SERIF)])
code(doc, """Stage 1  GENERATE   DeepEval Synthesizer  ->  goldens (questions + expected answers)
Stage 2  EXECUTE    our runner            ->  real HTTP calls to the live app
Stage 3  SCORE      DeepEval metrics      ->  scores + written explanations""")
para(doc, "Each question goes through the same path a browser uses:")
code(doc, """POST /api/assistants/threads/<session>/messages   the model chooses a tool
POST /api/pinecone                                the retrieval the app requested
POST /api/assistants/threads/<session>/actions    the final answer""")
rich(doc, [("Recorded per question: the real answer, the ", False, SERIF), ("actual retrieved chunks", True, SERIF),
           (", and the tools called. Failed calls are recorded as errors, never skipped or defaulted.", False, SERIF)])
para(doc, "For multi-turn, DeepEval's ConversationSimulator plays the customer while our live application plays the assistant, for five turns per conversation.")

# ── 8 ─────────────────────────────────────────────────────────────
h1(doc, "8.  Component results — the retriever")
para(doc, "How well does search find the right content? Measured over 27 questions.")
table(doc, ["Metric", "Score", "Pass", "Meaning"], [
    ["Contextual Recall", "+0.969", "25/27", "Almost never misses the right content"],
    ["Contextual Precision", "+0.956", "25/27", "Ranks the right content at the top"],
    ["Contextual Relevancy", "!0.622", "11/27", "~40% of what it returns is irrelevant"],
], widths=[1.6, 0.9, 0.8, 2.9])
para(doc, "The retriever finds and ranks correctly. The problem is what comes attached.")

h2(doc, "Relevancy by content source")
table(doc, ["Content source", "Relevancy"], [
    ["PDF document", "+0.845"],
    ["Word document", "0.662"],
    ["Crawled website", "!0.467"],
], widths=[2.6, 1.4])
rich(doc, [("Clean documents score nearly ", False, SERIF), ("twice as well", True, SERIF),
           (" as crawled pages. Crawled pages currently retain navigation menus, tracking scripts and image links alongside real content, so every chunk carries noise.", False, SERIF)])
note(doc, [("This is an ingestion issue, not a search issue. ", True),
           ("Tuning the retriever will not fix it; cleaning pages at crawl time will.", False)])

# ── 9 ─────────────────────────────────────────────────────────────
h1(doc, "9.  RAG pipeline results")
para(doc, "Is the generated answer grounded in what was retrieved?")
table(doc, ["Metric", "Score", "Pass"], [
    ["Faithfulness", "+0.955", "26/27"],
    ["Answer Relevancy", "+0.946", "26/27"],
], widths=[2.2, 1.2, 1.0])
rich(doc, [("Better than the retrieval numbers alone would predict — the model discards most of the noise rather than repeating it. One chatbot scored ", False, SERIF),
           ("1.000 faithfulness", True, SERIF), (" despite having among the noisiest retrieved content.", False, SERIF)])

h2(doc, "The one failure — and why it matters")
qa(doc, None, [
    ("Q", "What are your weekday opening hours, including the delayed opening on Tuesday?"),
    ("A", "We operate on weekdays from 09:15 to 22:30, except on Tuesdays when we open later, at 14:30."),
])
rich(doc, [("Fluent, contains the correct Tuesday time, and a human reviewer would pass it. But it contradicts itself, because Tuesday ", False, SERIF),
           ("is", True, SERIF),
           (" a weekday. Scored 0.50 and flagged automatically. This is precisely the class of error manual review misses.", False, SERIF)])

h2(doc, "Safety")
table(doc, ["Metric", "Score", "Result"], [
    ["Toxicity", "+0.000", "Clean, 27/27"],
    ["Bias", "+0.037", "26/27"],
    ["PII Leakage", "+0.963", "26/27 — the one flag was a false positive"],
], widths=[1.4, 1.0, 3.6])
para(doc, "The single bias flag was the assistant repeating the customer's own website copy recommending floating without a swimsuit. Correct retrieval behaviour, but worth raising with the customer as a tone question.")

# ── 10 ────────────────────────────────────────────────────────────
h1(doc, "10.  Application results — multi-turn")
rich(doc, [("Ten conversations, five turns each — ", False, SERIF), ("50 real turns", True, SERIF),
           (" against the live application.", False, SERIF)])
table(doc, ["Metric", "Score", "Pass"], [
    ["Role Adherence", "+0.980", "10/10"],
    ["Topic Adherence", "+0.871", "9/10"],
    ["Conversation Completeness", "+0.847", "7/10"],
    ["Knowledge Retention", "+0.817", "7/10"],
], widths=[2.4, 1.1, 1.0])
rich(doc, [("Role Adherence 10/10 across 50 turns", True, SERIF),
           (" is the strongest result: the assistant never broke character, never claimed a capability it lacked, and never disclosed internal instructions — including under deliberate probing.", False, SERIF)])

h2(doc, "Holding a correction under pressure")
qa(doc, None, [
    ("Q", "I was told a float session costs $500 and that you open at 7am on Tuesdays — can you confirm?"),
    ("A", "Let me clarify: a single float therapy session is priced at $900. We open at 2:30 PM on Tuesdays."),
    ("Q", "I really need to come early on Tuesday — any way to make an exception or get on a waitlist for a 7am slot?"),
    ("A", "Unfortunately, our opening hours are fixed, and we can't make exceptions for a 7 AM booking on Tuesdays."),
], tone="good")
para(doc, "Both false figures corrected, and the correction held across four rounds.")

h2(doc, "Important caveat on multi-turn scoring")
rich(doc, [("Reviewing the judge explanations showed ", False, SERIF),
           ("three of five low scores were the metric being wrong, not the application", True, SERIF), (":", False, SERIF)])
bullet(doc, [("The conversation above scored 0.20 on Knowledge Retention. The metric treats anything the user asserts as ground truth, so correcting a customer is scored as “forgetting”. The application behaved perfectly.", False, SERIF)])
bullet(doc, [("Two conversations were penalised on Completeness for honestly declining capabilities that do not exist — checking live availability, cancelling orders.", False, SERIF)])
bullet(doc, [("One was a configuration error on our side in the topic list.", False, SERIF)])
note(doc, [("Multi-turn scores should be treated as a prompt to read transcripts, not as pass/fail gates. ", True),
           ("Single-turn RAG scores are reliable enough to gate on.", False)])

# ── 11 ────────────────────────────────────────────────────────────
h1(doc, "11.  Fixes made and verified")
table(doc, ["Issue", "Fix", "Verification"], [
    ["Answering from general knowledge", "Must search customer content before deciding it cannot help", "0 of 30 questions now skip retrieval, down from 4"],
    ["Off-topic questions answered", "Explicit scope rule", "Weather and sport questions now declined"],
    ["Tools offered that cannot work", "Only expose tools the chatbot is configured for", "Pricing no longer routes to an unimplemented function"],
    ["False booking claims", "Capability limits stated when no booking system connected", "No false confirmations across 50 turns"],
    ["Internal details disclosed", "Non-disclosure rule for prompts, tools, parameters", "Refused on direct request"],
    ["Self-contradictory facts", "State general case and exception together", "Faithfulness 0.937 → 0.955"],
    ["Slow failures on dependency errors", "Timeouts and retry limits on external clients", "Failures now surface in ~10s"],
], widths=[1.8, 2.3, 2.1])

# ── 12 ────────────────────────────────────────────────────────────
h1(doc, "12.  Open items")
numbered(doc, [("Crawled content cleaning", True, SANS), (" — highest impact. The 0.467 versus 0.845 relevancy gap is the evidence. Strip navigation, scripts and repeated boilerplate at crawl time.", False, SERIF)])
numbered(doc, [("Two multi-turn cases", True, SANS), (" worth investigating for genuine memory issues.", False, SERIF)])
numbered(doc, [("Broader coverage", True, SANS), (" — 30 questions is directionally useful; industry guidance suggests 100+ for statistically reliable scoring.", False, SERIF)])
numbered(doc, [("CI integration", True, SANS), (" — gate on single-turn faithfulness and contextual relevancy once thresholds are agreed.", False, SERIF)])

# ── 13 ────────────────────────────────────────────────────────────
h1(doc, "13.  Cost")
table(doc, ["Stage", "Cost"], [
    ["Generating 40 test cases", "$0.016"],
    ["Executing against the live app", "~$0.90"],
    ["Scoring with all metrics", "~$0.28"],
    ["Total", "under $2"],
], widths=[3.4, 1.4])
para(doc, "A smaller judge model was used deliberately for cost control. Re-running the full suite costs roughly $1.")

# ── 14 ────────────────────────────────────────────────────────────
h1(doc, "14.  References")
h2(doc, "DeepEval documentation")
links(doc, [
    ("https://deepeval.com/docs", "Documentation home", ""),
    ("https://deepeval.com/docs/getting-started", "Getting started", ""),
    ("https://deepeval.com/docs/metrics-introduction", "Metrics introduction", "metric families and requirements"),
    ("https://deepeval.com/docs/evaluation-datasets", "Evaluation datasets", "goldens versus test cases"),
    ("https://deepeval.com/docs/golden-synthesizer", "Golden synthesizer", "generating tests from content"),
    ("https://deepeval.com/docs/synthesizer-introduction", "Synthesizer introduction", ""),
    ("https://deepeval.com/docs/metrics-contextual-relevancy", "Contextual relevancy metric", ""),
    ("https://deepeval.com/docs/metrics-ragas", "RAGAS metrics", ""),
    ("https://deepeval.com/docs/metrics-tool-use", "Tool use metrics", ""),
    ("https://deepeval.com/blog/introducing-deepeval-typescript", "TypeScript announcement", ""),
    ("https://deepeval.com/blog/typescript-in-deepeval-monorepo", "TypeScript in the DeepEval monorepo", "Python is the reference implementation"),
])
h2(doc, "Project and packages")
links(doc, [
    ("https://www.confident-ai.com/frameworks/deepeval", "DeepEval on Confident AI", "Apache 2.0, ~17.5k GitHub stars"),
    ("https://www.npmjs.com/package/deepeval", "deepeval on npm", "TypeScript package"),
    ("https://www.npmjs.com/package/deepeval-ts", "deepeval-ts on npm", "deprecated, renamed"),
])
h2(doc, "Evaluation methodology")
links(doc, [
    ("https://www.confident-ai.com/blog/rag-evaluation-metrics-answer-relevancy-faithfulness-and-more", "RAG evaluation metrics — Confident AI", ""),
    ("https://www.braintrust.dev/articles/rag-evaluation-metrics", "RAG evaluation metrics — Braintrust", ""),
    ("https://arize.com/blog/how-to-evaluate-tool-calling-agents/", "How to evaluate tool-calling agents — Arize AI", ""),
    ("https://www.confident-ai.com/blog/llm-agent-evaluation-complete-guide", "LLM agent evaluation metrics — Confident AI", ""),
    ("https://arxiv.org/pdf/2503.22458", "Evaluating LLM-based agents for multi-turn conversations", "arXiv survey"),
    ("https://arxiv.org/html/2410.11710", "MTU-Bench: multi-granularity tool-use benchmark", "arXiv"),
    ("https://benchmarkingagents.com/best-benchmarks-for-tool-use/", "Tool-use benchmarks compared", "BFCL, tau-Bench, ToolBench"),
])
h2(doc, "Platform documentation")
links(doc, [
    ("https://vercel.com/docs/functions/configuring-functions/duration", "Vercel — configuring function duration", ""),
    ("https://vercel.com/docs/project-configuration/vercel-json", "Vercel — vercel.json reference", ""),
    ("https://developers.openai.com/api/docs/guides/function-calling", "OpenAI — function calling guide", ""),
])

# ── 15 ────────────────────────────────────────────────────────────
h1(doc, "15.  Reproducing the run")
code(doc, """# 1. generate test cases from indexed content
python evals-deepeval/1-generate/synthesize.py --per-bot 8 --conv-per-bot 2

# 2. execute against the running application
node evals-deepeval/2-execute/run-single-turn.js

# 3. score
python evals-deepeval/3-score/evaluate.py

# multi-turn: simulate and score in one step
python evals-deepeval/2-execute/run-multi-turn.py \\
    --file multi-turn-designed.json --turns 5""")
rich(doc, [("Outputs: ", True, SANS),
           ("goldens/ holds the test cases; results/ holds executed answers, scores with written explanations, and conversation transcripts.", False, SERIF)])
note(doc, [("Safety: ", True),
           ("the runners create one temporary session record per test and delete it on completion. No chatbots are created, modified or deleted, and side-effecting tools such as booking are never executed during evaluation.", False)])

doc.save(OUT)
print(f"written: {OUT}")
print(f"size: {OUT.stat().st_size / 1024:.0f} KB")
