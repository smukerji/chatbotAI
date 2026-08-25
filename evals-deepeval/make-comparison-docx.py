"""
Builds the Word version of the platform comparison report.

Same helper set as make-docx.py - python-docx has no public hyperlink API, so
add_hyperlink() writes the relationship and w:hyperlink element directly, which
is what makes the source links actually clickable in Word.

  evals-deepeval/.venv/Scripts/python.exe evals-deepeval/make-comparison-docx.py
"""

import pathlib

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

OUT = pathlib.Path(__file__).resolve().parent / "Eval-Platform-Comparison.docx"

# palette mirrors platform-comparison.html
INK = RGBColor(0x10, 0x1A, 0x1F)
ACCENT = RGBColor(0x14, 0x59, 0x5E)
MUTED = RGBColor(0x5E, 0x70, 0x77)
GOOD = RGBColor(0x2C, 0x6E, 0x52)
POOR = RGBColor(0x8C, 0x2F, 0x2F)
WARN = RGBColor(0x87, 0x56, 0x0A)

SANS = "IBM Plex Sans Condensed"
SERIF = "IBM Plex Serif"
MONO = "IBM Plex Mono"


def set_font(run, name, size=None, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size) if size else run.font.size
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), name)


def add_hyperlink(paragraph, url, text, size=10.5):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")

    fonts = OxmlElement("w:rFonts")
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        fonts.set(qn(attr), SANS)
    rpr.append(fonts)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "14595E")
    rpr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    rpr.append(sz)

    run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    link.append(run)
    paragraph._p.append(link)
    return link


def shade(cell, hex_color):
    tcpr = cell._tc.get_or_add_tcPr()
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hex_color)
    tcpr.append(el)


def para(doc, text="", font=SERIF, size=10.5, bold=False, italic=False,
         color=INK, space_after=8, space_before=0, indent=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    if text:
        set_font(p.add_run(text), font, size, bold, italic, color)
    return p


def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(9)
    set_font(p.add_run(text), SANS, 15, True, False, INK)
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), "14595E")
    pbdr.append(bottom)
    p._p.get_or_add_pPr().append(pbdr)
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    set_font(p.add_run(text), SANS, 11.5, True, False, INK)
    return p


def bullet(doc, parts, size=10.5):
    if isinstance(parts, str):
        parts = [(parts, False, SERIF)]
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    for text, bold, font in parts:
        set_font(p.add_run(text), font, size, bold, False, INK)
    return p


def numbered(doc, parts, size=10.5):
    if isinstance(parts, str):
        parts = [(parts, False, SERIF)]
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    for text, bold, font in parts:
        set_font(p.add_run(text), font, size, bold, False, INK)
    return p


def caption(doc, text):
    para(doc, text, MONO, 8, True, False, MUTED, space_after=3, space_before=6)


def table(doc, headers, rows, widths=None, highlight=None):
    """highlight = index of the row to shade as the recommended option."""
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        set_font(p.add_run(htext.upper()), MONO, 8, True, False, MUTED)
        shade(hdr[i], "E7EEEF")
    for r_idx, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            if highlight is not None and r_idx == highlight:
                shade(cells[i], "E4F0F0")
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            colour, mono = INK, False
            if isinstance(val, str) and val[:1] in "!+~":
                marker, val = val[0], val[1:]
                colour = {"!": POOR, "+": GOOD, "~": WARN}[marker]
                mono = True
            elif isinstance(val, str) and val.startswith("$"):
                mono = True
            set_font(p.add_run(str(val)), MONO if mono else SANS, 9.5, False, False, colour)
    if widths:
        for i, w in enumerate(widths):
            for r in t.rows:
                r.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def note(doc, label, text, fill="F2F6F6", label_color=WARN):
    t = doc.add_table(rows=1, cols=1)
    cell = t.rows[0].cells[0]
    cell.text = ""
    shade(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    set_font(p.add_run(label.upper()), MONO, 8, True, False, label_color)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(2)
    set_font(p2.add_run(text), SERIF, 10.5, False, False, INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def links(doc, items):
    for url, text, desc in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Inches(0.16)
        if desc:
            set_font(p.add_run(f"{desc}  "), SANS, 9.5, False, False, MUTED)
        add_hyperlink(p, url, text)


# ══════════════════════════════════════════════════════════════════
doc = Document()

sec = doc.sections[0]
sec.left_margin = sec.right_margin = Inches(1.05)
sec.top_margin = Inches(0.95)
sec.bottom_margin = Inches(0.95)

normal = doc.styles["Normal"]
normal.font.name = SERIF
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.2

# ── cover ─────────────────────────────────────────────────────────
para(doc, "EVALUATION PROGRAMME  ·  PLATFORM REVIEW", MONO, 9, True, False, ACCENT, space_after=10)
para(doc, "Choosing Our Eval Platform", SANS, 25, True, False, INK, space_after=6)
para(doc,
     "A comparison of three platforms for running evaluation on live production traffic and "
     "surfacing the results in the product: Confident AI, LangSmith and Langfuse.",
     SERIF, 12.5, False, True, MUTED, space_after=16)

table(doc, ["Field", "Detail"], [
    ["Date", "21 August 2026"],
    ["Platforms compared", "Confident AI, LangSmith, Langfuse"],
    ["Recommendation", "Langfuse, on the $29/month cloud plan"],
    ["Verification", "Pricing and feature details taken from vendor pages on 21 Aug 2026"],
], widths=[1.5, 4.7])

# ── 1 ─────────────────────────────────────────────────────────────
h1(doc, "1.  Scope")

bullet(doc, [("Evaluation completed so far is ", False, SERIF), ("offline", True, SERIF),
             (": a generated test set of questions and multi-turn conversations, run against the "
              "chatbots and scored.", False, SERIF)])
bullet(doc, [("The next stage is ", False, SERIF), ("online evaluation", True, SERIF),
             (": scoring a sample of live customer conversations on a continuous basis.", False, SERIF)])
bullet(doc, "Online evaluation requires a platform that traces the running application and stores the scores.")
bullet(doc, "This document compares the three candidates and recommends one.")

caption(doc, "Offline and online evaluation compared")
table(doc, ["", "Offline", "Online"], [
    ["Input", "Generated test set", "Live customer conversations"],
    ["Frequency", "On demand", "Continuous, sampled"],
    ["Requires", "A test suite", "Tracing in the live application"],
    ["Status", "+Complete", "!Not started"],
], widths=[1.1, 2.5, 2.6])

# ── 2 ─────────────────────────────────────────────────────────────
h1(doc, "2.  Requirements")

para(doc, "Criteria used for the comparison:")
numbered(doc, "Online evaluation on production traffic, with sampling to control cost.")
numbered(doc, "Scores readable through an API, so the admin view and the customer evaluation tab "
              "can be rendered inside our own product.")
numbered(doc, "No cap on the number of chatbots or projects.")
numbered(doc, "Pricing that does not scale per user seat.")
numbered(doc, "Self-hosting available, in case a customer requires conversation data to stay on "
              "our infrastructure.")
numbered(doc, "A supported TypeScript SDK, since the application is a TypeScript Next.js codebase.")

# ── 3 ─────────────────────────────────────────────────────────────
h1(doc, "3.  Confident AI")

para(doc, "Hosted platform built by the makers of DeepEval, the framework used for our offline suite.")

h2(doc, "Features")
bullet(doc, "Tracing of each execution as traces, spans and threads.")
bullet(doc, [("Online evaluation: ", False, SERIF),
             ('"continuously evaluate production traffic using 50+ metrics"', False, SERIF),
             (", run in real time as traces are ingested or retrospectively.", False, SERIF)])
bullet(doc, "Sampling controls, monitors and alerting on quality, latency and cost.")
bullet(doc, "Datasets, prompt versioning, human annotation, CI/CD integration.")
bullet(doc, "Instrumentation via an @observe decorator, framework integrations, or OpenTelemetry. "
            "Python and TypeScript supported.")
bullet(doc, [('REST API: "Every part of Confident AI is exposed as an API".', False, SERIF)])
bullet(doc, "No self-hosted deployment option.")

h2(doc, "Pricing")
caption(doc, "Confident AI · verified 21 Aug 2026")
table(doc, ["Plan", "Price / mo", "Trace spans", "Projects", "Seats", "Online evals"], [
    ["Free", "$0", "1 GB-month", "1", "2", "!Not included"],
    ["Starter", "$200", "5 GB-months", "5", "Unlimited", "+Included"],
    ["Team", "$2,000", "75 GB-months", "Unlimited", "Unlimited", "+Included"],
    ["Enterprise", "Custom", "Unlimited", "Unlimited", "Unlimited", "+Unlimited"],
], widths=[0.85, 0.85, 1.1, 0.9, 0.9, 1.6])

para(doc, "Overage: $1 per GB-month ingested or retained. Judge-model tokens billed separately at "
          "~$0.05/M input tokens and $0.40/M output tokens on both paid tiers.",
     SERIF, 9.5, False, False, MUTED)

h2(doc, "Assessment against the criteria")
bullet(doc, [("Online evaluation is not available on the Free plan. Entry price is ", False, SERIF),
             ("$200/month", True, SERIF), (".", False, SERIF)])
bullet(doc, [("The Starter plan is limited to ", False, SERIF), ("5 projects", True, SERIF),
             (". Project is the isolation boundary between customers.", False, SERIF)])
bullet(doc, [("Unlimited projects begin at the Team plan, ", False, SERIF),
             ("$2,000/month", True, SERIF), (".", False, SERIF)])
bullet(doc, "No self-hosting, so conversation data cannot be kept on our infrastructure.")
bullet(doc, "API access and TypeScript support: both available.")

# ── 4 ─────────────────────────────────────────────────────────────
h1(doc, "4.  LangSmith")

para(doc, "LangChain's hosted observability and evaluation platform. Proprietary; the client SDKs "
          "are open source, the platform is not.")

h2(doc, "Features")
bullet(doc, "Tracing, datasets, experiments, prompt management and online evaluators on production traces.")
bullet(doc, "Closest integration with LangChain and LangGraph; tracing for those frameworks enables "
            "via a single environment variable.")
bullet(doc, "Python and JS/TS SDKs. OpenTelemetry ingestion supported as a secondary path.")
bullet(doc, "Self-hosted and hybrid deployment listed on the Enterprise plan only.")

h2(doc, "Pricing")
caption(doc, "LangSmith · verified 21 Aug 2026")
table(doc, ["Plan", "Price", "Seats", "Included traces", "Self-hosting"], [
    ["Developer", "$0 / seat / mo", "1", "5k base traces / mo", "!No"],
    ["Plus", "$39 / seat / mo", "Billed per seat", "10k base traces / mo", "!No"],
    ["Enterprise", "Custom", "Custom", "Custom", "+Included"],
], widths=[1.0, 1.25, 1.25, 1.4, 1.3])

para(doc, "Usage beyond the allowance is metered as $1.50 / LCU (compute) and $1.00 / LSU (storage) "
          "following the July 2026 repricing. Base traces are retained 14 days; extended traces 400 days.",
     SERIF, 9.5, False, False, MUTED)

h2(doc, "Assessment against the criteria")
bullet(doc, "Pricing scales on two axes: per seat and per trace.")
bullet(doc, "Self-hosting requires an Enterprise contract.")
bullet(doc, "Its principal advantage is depth of LangChain and LangGraph integration. Our "
            "conversation path calls the OpenAI Responses API directly, so that advantage applies "
            "only partially.")
bullet(doc, "API access and TypeScript support: both available.")

# ── 5 ─────────────────────────────────────────────────────────────
h1(doc, "5.  Langfuse")

para(doc, "Open-source observability and evaluation platform, MIT licensed. All product features "
          "were moved to MIT in June 2025; SCIM, audit logs, project-level access control and UI "
          "customisation remain commercial. Acquired by ClickHouse in January 2026, licence unchanged.")

h2(doc, "Features")
bullet(doc, [("Tracing", True, SERIF),
             (' — "all LLM and non-LLM calls, including retrieval, embedding, API calls", with '
              "sessions, user tracking and agent-graph visualisation.", False, SERIF)])
bullet(doc, [("Evaluation", True, SERIF),
             (" — LLM-as-a-judge, code evaluators, user feedback, manual labelling, custom "
              "pipelines. Open source since v3.65.0.", False, SERIF)])
bullet(doc, [("Code evaluators", True, SERIF),
             (" — TypeScript or Python evaluate functions written in the UI, shipped May 2026. "
              "Deterministic checks run without a judge model call.", False, SERIF)])
bullet(doc, [("Sampling", True, SERIF),
             (' — "configure sampling percentage (e.g. 5%) to manage evaluation costs and '
              'throughput", applied deterministically per observation.', False, SERIF)])
bullet(doc, [("Prompt management", True, SERIF), (" — versioning and label-based deployment.", False, SERIF)])
bullet(doc, [("Datasets and annotation queues", True, SERIF),
             (" — existing offline goldens and human review.", False, SERIF)])
bullet(doc, [("Public API", True, SERIF),
             (" — scores, observations, metrics and dashboard endpoints, typed in the JS/TS SDK.", False, SERIF)])
bullet(doc, [("TypeScript SDK v5", True, SERIF),
             (' — OpenTelemetry-based. Documented Next.js path: "register the LangfuseSpanProcessor '
              'via registerOTel from @vercel/otel as long as you are on @vercel/otel v2 or later".', False, SERIF)])
bullet(doc, [("Self-hosting", True, SERIF),
             (" — free, with all core platform features and APIs.", False, SERIF)])

h2(doc, "Pricing")
caption(doc, "Langfuse · verified 21 Aug 2026")
table(doc, ["Plan", "Price / mo", "Included units", "Retention", "Seats", "Projects"], [
    ["Hobby", "$0", "50,000", "30 days", "2", "Unlimited"],
    ["Core", "$29", "100,000", "90 days", "Unlimited", "Unlimited"],
    ["Pro", "$199", "100,000", "3 years", "Unlimited", "Unlimited"],
    ["Enterprise", "$2,499", "100,000", "3 years", "Unlimited", "Unlimited"],
    ["Self-hosted", "$0", "Unlimited", "Our choice", "Unlimited", "Unlimited"],
], widths=[1.0, 0.85, 1.15, 1.0, 1.1, 1.1], highlight=1)

para(doc, 'A billable unit is "any tracing data point sent to the platform — including traces, '
          'observations, and scores". Overage is $8 per 100k units, on a graduated scale down to '
          "$6 per 100k above 50M units. No plan charges per seat.",
     SERIF, 9.5, False, False, MUTED)

note(doc, "Self-hosting requirement",
     "Self-hosted Langfuse carries no licence cost but requires PostgreSQL, ClickHouse, Redis and "
     "object storage to be operated by us. Because the licence is MIT and the API is identical "
     "between cloud and self-hosted, moving from cloud to self-hosted later does not require "
     "application changes.")

# ── 6 ─────────────────────────────────────────────────────────────
h1(doc, "6.  Side by side")

caption(doc, "All three platforms against the criteria in section 2")
table(doc, ["Criterion", "Langfuse", "LangSmith", "Confident AI"], [
    ["Entry price, online evals", "$0 self-hosted / $29 cloud", "$39 / seat / mo", "$200 / mo"],
    ["Per-seat charge", "+None", "!$39 per seat", "+None"],
    ["Project cap", "+None", "+None", "~5 on Starter"],
    ["Self-hosting", "+Free", "~Enterprise only", "!Not offered"],
    ["Licence", "MIT, full platform", "Proprietary", "Proprietary"],
    ["API for our own dashboards", "+Yes", "+Yes", "+Yes"],
    ["TypeScript SDK", "v5, OTel-native", "Yes, OTel secondary", "Yes"],
    ["Sampling", "+Yes", "+Yes", "+Yes"],
], widths=[1.6, 1.75, 1.5, 1.35])

# ── 7 ─────────────────────────────────────────────────────────────
h1(doc, "7.  Recommendation")

note(doc, "Recommendation", "Langfuse, on the $29/month cloud plan.",
     fill="E4F0F0", label_color=ACCENT)

bullet(doc, "Lowest entry price for online evaluation of the three.")
bullet(doc, "No per-seat charge and no project cap, so cost does not move with team size or chatbot count.")
bullet(doc, "Self-hosting available at no licence cost if a customer requires data residency, with "
            "no application changes needed to move.")
bullet(doc, "MIT licence for the full platform.")
bullet(doc, "TypeScript SDK is OpenTelemetry-native with a documented Next.js integration path.")
bullet(doc, "Scores are readable through the public API, so the admin view and customer evaluation "
            "tab are built inside our own product.")

para(doc, "Judge-model tokens are charged by the model provider on top of the platform cost, on "
          "whichever platform is chosen. Our completed offline run — 30 single-turn questions and "
          "10 multi-turn conversations, scored on a small judge model — cost under $2 in total. "
          "Sampling percentage and the use of code evaluators, which require no model call, both "
          "reduce this.", space_before=6)

# ── 8 ─────────────────────────────────────────────────────────────
h1(doc, "8.  Prompt-injection measure")

bullet(doc, "DeepEval's core metric set does not include a prompt-injection metric.")
bullet(doc, "None of the three platforms above provides one as a standard metric.")
bullet(doc, "Injection testing is adversarial — attacks are constructed and checked for success — "
            "so it runs as a pre-release check rather than as a sample of live traffic.")
bullet(doc, "promptfoo's red-team module generates adversarial inputs across prompt injection, "
            "jailbreaks, PII leakage and hijacking, and reports which attacks succeeded. Open source.")
bullet(doc, "DeepTeam, from the DeepEval authors, is the alternative and shares the vocabulary of "
            "our existing suite.")

# ── 9 ─────────────────────────────────────────────────────────────
h1(doc, "9.  Next steps")

numbered(doc, [("Instrument the application. ", True, SERIF),
               ("Add the Langfuse TypeScript SDK to the chat path so conversations, retrieval steps "
                "and tool calls are traced.", False, SERIF)])
numbered(doc, [("Enable sampled online evaluation. ", True, SERIF),
               ("Configure the judge model and code evaluators, start at a low sampling percentage, "
                "and compare results against the existing offline scores.", False, SERIF)])
numbered(doc, [("Build the admin view. ", True, SERIF),
               ("Read scores through the Metrics API into an internal page covering all chatbots.", False, SERIF)])
numbered(doc, [("Build the customer evaluation tab. ", True, SERIF),
               ("The same data filtered to a single account.", False, SERIF)])
numbered(doc, [("Add the prompt-injection check. ", True, SERIF),
               ("promptfoo red-team run wired into the release process.", False, SERIF)])

para(doc, "Each step is usable on its own, so the sequence can be paused or re-ordered without "
          "discarding work.", space_before=6)

# ── 10 ────────────────────────────────────────────────────────────
h1(doc, "10.  Sources")

para(doc, "All figures and quotations are taken from the vendors' published pages, checked on "
          "21 August 2026.")

links(doc, [
    ("https://langfuse.com/pricing", "langfuse.com/pricing", "Langfuse · cloud pricing"),
    ("https://langfuse.com/pricing-self-host", "langfuse.com/pricing-self-host", "Langfuse · self-host pricing"),
    ("https://langfuse.com/docs/observability/sdk/typescript/overview", "SDK overview and Next.js setup", "Langfuse · TypeScript SDK"),
    ("https://langfuse.com/docs/evaluation/evaluation-methods/llm-as-a-judge", "Evaluators and sampling", "Langfuse · LLM-as-a-judge"),
    ("https://langfuse.com/docs/api-and-data-platform/features/public-api", "REST API reference", "Langfuse · public API"),
    ("https://langfuse.com/docs/metrics/features/metrics-api", "Querying scores", "Langfuse · metrics API"),
    ("https://www.langchain.com/pricing-langsmith", "langchain.com/pricing-langsmith", "LangSmith · pricing"),
    ("https://www.confident-ai.com/pricing", "confident-ai.com/pricing", "Confident AI · pricing"),
    ("https://www.confident-ai.com/docs/llm-tracing/introduction", "LLM tracing introduction", "Confident AI · tracing"),
    ("https://deepeval.com/docs", "deepeval.com/docs", "DeepEval · documentation"),
    ("https://www.promptfoo.dev/docs/red-team/", "Adversarial testing", "promptfoo · red teaming"),
    ("https://www.datacamp.com/blog/langfuse-vs-langsmith", "DataCamp · Langfuse vs LangSmith", "Independent comparison"),
])

para(doc, "")
para(doc, "21 August 2026 · pricing subject to vendor change; re-verify before contracting",
     MONO, 8.5, False, False, MUTED)

doc.save(OUT)
print(f"-> {OUT}")
